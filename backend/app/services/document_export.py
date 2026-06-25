"""
DOCX and PDF rendering for AI-drafted documents.

The documents arrive as plain Markdown text. We parse a deliberately small
subset (headings, bold/italic, bullets, blank-line paragraphs) and emit a
uniform block model that both renderers consume — no system dependencies
(LibreOffice/Cairo) so the Docker image stays slim.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Literal


BlockKind = Literal["heading", "paragraph", "bullets"]


@dataclass
class Run:
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class Block:
    kind: BlockKind
    level: int = 0  # heading level, 1..6
    runs: list[Run] = field(default_factory=list)
    items: list[list[Run]] = field(default_factory=list)  # only for "bullets"


# Match either **bold** or *italic*; longest first.
_INLINE_RE = re.compile(r"(\*\*([^\*]+)\*\*|\*([^\*]+)\*)")


def _parse_runs(text: str) -> list[Run]:
    runs: list[Run] = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            runs.append(Run(text[pos:m.start()]))
        if m.group(2) is not None:
            runs.append(Run(m.group(2), bold=True))
        elif m.group(3) is not None:
            runs.append(Run(m.group(3), italic=True))
        pos = m.end()
    if pos < len(text):
        runs.append(Run(text[pos:]))
    if not runs:
        runs.append(Run(text))
    return runs


def parse_markdown(text: str) -> list[Block]:
    """Tiny markdown subset → block list."""
    lines = text.replace("\r\n", "\n").split("\n")
    blocks: list[Block] = []
    paragraph_buf: list[str] = []
    bullet_buf: list[str] = []

    def flush_paragraph():
        if paragraph_buf:
            text = " ".join(line.strip() for line in paragraph_buf).strip()
            if text:
                blocks.append(Block(kind="paragraph", runs=_parse_runs(text)))
            paragraph_buf.clear()

    def flush_bullets():
        if bullet_buf:
            items = [_parse_runs(item) for item in bullet_buf]
            blocks.append(Block(kind="bullets", items=items))
            bullet_buf.clear()

    for raw in lines:
        line = raw.rstrip()

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            flush_paragraph()
            flush_bullets()
            level = len(m.group(1))
            blocks.append(Block(kind="heading", level=level, runs=_parse_runs(m.group(2))))
            continue

        # Bullet
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            flush_paragraph()
            bullet_buf.append(m.group(1))
            continue

        # Blank line → flush
        if not line.strip():
            flush_paragraph()
            flush_bullets()
            continue

        # Anything else accumulates into the current paragraph
        flush_bullets()
        paragraph_buf.append(line)

    flush_paragraph()
    flush_bullets()
    return blocks


# ----- DOCX -----

def to_docx(content: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    # Sensible legal-document defaults
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.size = Pt(16)

    for block in parse_markdown(content):
        if block.kind == "heading":
            p = doc.add_heading("", level=min(block.level, 4))
            for r in block.runs:
                run = p.add_run(r.text)
                run.bold = r.bold or True  # headings already bold by style; explicit is fine
                run.italic = r.italic
        elif block.kind == "bullets":
            for item_runs in block.items:
                p = doc.add_paragraph(style="List Bullet")
                for r in item_runs:
                    run = p.add_run(r.text)
                    run.bold = r.bold
                    run.italic = r.italic
        else:  # paragraph
            p = doc.add_paragraph()
            for r in block.runs:
                run = p.add_run(r.text)
                run.bold = r.bold
                run.italic = r.italic

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----- PDF -----

def to_pdf(content: str, title: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title", parent=styles["Title"], fontSize=18, spaceAfter=18, alignment=1,
    )
    body_style = ParagraphStyle(
        "body", parent=styles["BodyText"], fontSize=11, leading=15,
        alignment=TA_JUSTIFY, spaceAfter=8,
    )
    heading_styles = {
        i: ParagraphStyle(
            f"h{i}", parent=styles[f"Heading{min(i, 4)}"],
            fontSize=max(11, 18 - 2 * i), spaceBefore=14, spaceAfter=8,
        )
        for i in range(1, 7)
    }

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=title,
    )

    def runs_to_html(runs: list[Run]) -> str:
        out = []
        for r in runs:
            txt = (
                r.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            if r.bold and r.italic:
                out.append(f"<b><i>{txt}</i></b>")
            elif r.bold:
                out.append(f"<b>{txt}</b>")
            elif r.italic:
                out.append(f"<i>{txt}</i>")
            else:
                out.append(txt)
        return "".join(out)

    story = []
    if title:
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3 * cm))

    for block in parse_markdown(content):
        if block.kind == "heading":
            story.append(Paragraph(runs_to_html(block.runs), heading_styles[block.level]))
        elif block.kind == "bullets":
            items = [
                ListItem(Paragraph(runs_to_html(item), body_style), leftIndent=12)
                for item in block.items
            ]
            story.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=15))
            story.append(Spacer(1, 0.2 * cm))
        else:
            story.append(Paragraph(runs_to_html(block.runs), body_style))

    doc.build(story)
    return buf.getvalue()
